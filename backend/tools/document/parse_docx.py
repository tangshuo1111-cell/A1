"""
V16 R1：docx 文档工具（MCP-compatible Adapter）。

技术路线（按技术方案 CSV 第 5 行）：
  - MarkItDown 主入口（microsoft/markitdown）：统一输出 Markdown 文本
  - python-docx 表格 / 段落兜底：当 MarkItDown 未安装或失败时使用
  - 两者都未安装 → error_code=dependency_missing，不静默降级

metadata 至少包含：
  source_type=docx / filename / heading_level / paragraph_index /
  table_index / paragraph_count / table_count / parser_name /
  extract_method / content_hash
"""

from __future__ import annotations

import hashlib
import time
from pathlib import Path
from typing import Any

from tools.document.errors import (
    DEPENDENCY_MISSING,
    DOCX_NO_CONTENT,
    EMPTY_EXTRACTED_TEXT,
    FILE_NOT_FOUND,
    FILE_TOO_LARGE,
    PARSE_FAILED,
    UNSUPPORTED_FILE_TYPE,
)
from tools.document.limits import MAX_DOCX_PARAGRAPHS, MAX_FILE_BYTES
from tools.document.quality import assess_quality
from tools.document.registry import DocumentToolSchema, register
from tools.document.tool_result import DocumentToolResult

_INPUT_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "file_path": {"type": "string"},
        "file_content": {"type": ["string", "object", "null"], "description": "bytes 或 None"},
    },
    "required": ["file_path"],
}

_OUTPUT_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "status": {"type": "string"},
        "text": {"type": "string"},
        "metadata": {"type": "object"},
        "quality": {"type": "object"},
        "structured_data": {"type": "object"},
        "error_code": {"type": "string"},
    },
}


def _read_bytes(
    file_path: str, file_content: bytes | None
) -> tuple[bytes | None, str]:
    """返回 (content_bytes, error_code)。"""
    if file_content is not None:
        if len(file_content) > MAX_FILE_BYTES:
            return None, FILE_TOO_LARGE
        return file_content, ""
    p = Path(file_path)
    if not p.exists():
        return None, FILE_NOT_FOUND
    raw = p.read_bytes()
    if len(raw) > MAX_FILE_BYTES:
        return None, FILE_TOO_LARGE
    return raw, ""


def _parse_with_markitdown(
    file_path: str, content_bytes: bytes
) -> tuple[str, str, str]:
    """
    尝试用 MarkItDown 解析 docx。
    返回 (text, extract_method, error_msg)；成功时 error_msg 为空。
    """
    try:
        from markitdown import MarkItDown
    except ImportError:
        return "", "markitdown", "MarkItDown not installed"

    import os
    import tempfile

    try:
        # MarkItDown 支持文件路径或 file-like，用临时文件最稳
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content_bytes)
            tmp_path = tmp.name
        try:
            md = MarkItDown()
            result = md.convert(tmp_path)
            text = (result.text_content or "").strip()
            return text, "markitdown", ""
        finally:
            try:  # noqa: SIM105
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception as e:  # noqa: BLE001 - 解析失败统一返回错误串供上游降级
        return "", "markitdown", str(e)


def _parse_with_python_docx(
    content_bytes: bytes,
) -> tuple[str, dict[str, Any], str]:
    """
    用 python-docx 解析 docx，提取段落、标题、表格。
    返回 (text, structured_data, error_msg)。
    """
    try:
        import docx as _docx  # python-docx
    except ImportError:
        return "", {}, "python-docx not installed"

    import io
    try:
        doc = _docx.Document(io.BytesIO(content_bytes))
    except (OSError, ValueError, RuntimeError) as e:
        return "", {}, f"python-docx open failed: {e}"

    parts: list[str] = []
    paragraphs_meta: list[dict] = []
    tables_meta: list[dict] = []

    para_index = 0
    for para in doc.paragraphs[:MAX_DOCX_PARAGRAPHS]:
        t = para.text.strip()
        if not t:
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        heading_level = 0
        if "heading" in style_name:
            try:
                heading_level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                heading_level = 1
        paragraphs_meta.append({
            "paragraph_index": para_index,
            "heading_level": heading_level,
            "text_preview": t[:100],
        })
        if heading_level > 0:
            parts.append(f"{'#' * heading_level} {t}")
        else:
            parts.append(t)
        para_index += 1

    table_index = 0
    for table in doc.tables:
        table_text_parts: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_text_parts.append(" | ".join(cells))
        table_text = "\n".join(table_text_parts)
        if table_text.strip():
            tables_meta.append({
                "table_index": table_index,
                "row_count": len(table.rows),
                "col_count": len(table.columns) if table.rows else 0,
                "text_preview": table_text[:200],
            })
            parts.append(f"[表格{table_index}]\n{table_text}")
        table_index += 1

    text = "\n\n".join(parts).strip()
    structured_data = {
        "paragraphs": paragraphs_meta,
        "tables": tables_meta,
        "paragraph_count": para_index,
        "table_count": table_index,
    }
    return text, structured_data, ""


def _parse_docx(
    file_path: str,
    file_content: bytes | None = None,
) -> DocumentToolResult:
    """docx 文档解析 MCP-compatible Adapter。"""
    t0 = time.monotonic()
    path = Path(file_path)
    file_name = path.name or "unknown"
    trace: list[str] = [f"v16:parse_docx:start file={file_name}"]

    if path.suffix.lower() not in (".docx",):
        return DocumentToolResult(
            tool_name="parse_docx",
            source_type="docx",
            status="failed",
            error_code=UNSUPPORTED_FILE_TYPE,
            failure_reason=f"不支持的文件类型 {path.suffix}，parse_docx 只接受 .docx",
            trace=trace,
        )

    content_bytes, err = _read_bytes(file_path, file_content)
    if err:
        return DocumentToolResult(
            tool_name="parse_docx",
            source_type="docx",
            status="failed",
            error_code=err,
            failure_reason=f"文件读取失败: {err} ({file_path})",
            next_action_hint="请检查文件是否存在/是否过大",
            trace=trace,
        )

    assert content_bytes is not None

    # 1. 先尝试 MarkItDown
    text, method, md_err = _parse_with_markitdown(file_path, content_bytes)
    structured_data: dict[str, Any] = {}

    if text:
        trace.append(f"v16:parse_docx:markitdown=ok text_length={len(text)}")
        extract_method = "markitdown"
        # MarkItDown 不提供结构化段落/表格信息，补 python-docx structured_data
        _, sd, _ = _parse_with_python_docx(content_bytes)
        structured_data = sd
    else:
        # 2. MarkItDown 失败 → 尝试 python-docx
        if md_err:
            trace.append(f"v16:parse_docx:markitdown_failed reason={md_err[:80]}")
        text2, sd2, docx_err = _parse_with_python_docx(content_bytes)
        if text2:
            text = text2
            structured_data = sd2
            extract_method = "python_docx"
            trace.append(f"v16:parse_docx:python_docx=ok text_length={len(text)}")
        else:
            # 两者都失败 → 判断依赖缺失还是内容为空
            if "not installed" in md_err and "not installed" in docx_err:
                return DocumentToolResult(
                    tool_name="parse_docx",
                    source_type="docx",
                    status="failed",
                    error_code=DEPENDENCY_MISSING,
                    failure_reason="MarkItDown 和 python-docx 均未安装",
                    next_action_hint="请执行 pip install markitdown python-docx",
                    trace=trace,
                )
            if docx_err:
                return DocumentToolResult(
                    tool_name="parse_docx",
                    source_type="docx",
                    status="failed",
                    error_code=PARSE_FAILED,
                    failure_reason=f"docx 解析失败: {docx_err}",
                    trace=trace,
                )
            # 有库但内容为空
            if (
                structured_data.get("paragraph_count", 0) == 0
                and structured_data.get("table_count", 0) == 0
            ):
                return DocumentToolResult(
                    tool_name="parse_docx",
                    source_type="docx",
                    status="failed",
                    error_code=DOCX_NO_CONTENT,
                    failure_reason="docx 文件无有效段落和表格",
                    trace=trace,
                )
            return DocumentToolResult(
                tool_name="parse_docx",
                source_type="docx",
                status="failed",
                error_code=EMPTY_EXTRACTED_TEXT,
                failure_reason="docx 解析后内容为空",
                trace=trace,
            )

    if not text.strip():
        return DocumentToolResult(
            tool_name="parse_docx",
            source_type="docx",
            status="failed",
            error_code=EMPTY_EXTRACTED_TEXT,
            failure_reason="docx 提取内容为空",
            trace=trace,
        )

    quality = assess_quality(text, "docx")
    content_hash = quality.get("content_hash", hashlib.md5(text.encode()).hexdigest())

    metadata: dict[str, Any] = {
        "source_type": "docx",
        "filename": file_name,
        "heading_level": None,  # 在 structured_data.paragraphs 里
        "paragraph_count": structured_data.get("paragraph_count", 0),
        "table_count": structured_data.get("table_count", 0),
        "paragraph_index": 0,
        "table_index": 0,
        "parser_name": "parse_docx",
        "extract_method": extract_method,
        "content_hash": content_hash,
    }

    duration_ms = (time.monotonic() - t0) * 1000
    trace.append(
        f"v16:parse_docx:done method={extract_method} "
        f"text_length={len(text)} quality={quality.get('quality_level')}"
    )

    return DocumentToolResult(
        tool_name="parse_docx",
        source_type="docx",
        status="success",
        text=text,
        structured_data=structured_data,
        metadata=metadata,
        quality=quality,
        warnings=quality.get("warnings", []),
        duration_ms=duration_ms,
        trace=trace,
    )


register(DocumentToolSchema(
    tool_name="parse_docx",
    description="解析 .docx Word 文件，提取段落、标题、表格，返回 ToolResult",
    input_schema=_INPUT_SCHEMA,
    output_schema=_OUTPUT_SCHEMA,
    call_fn=_parse_docx,
))
