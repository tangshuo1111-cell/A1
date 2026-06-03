"""
V16 R1：测试 fixtures 动态生成工厂。

运行：python tests/_fixtures/v16_doc_factory.py
     或在测试中调用 ensure_fixtures() 确保文件存在。

生成到固定目录：tests/fixtures/v16_materials/
每类样本说明见下方 FIXTURES_MANIFEST。
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

# ── 固定输出目录 ──────────────────────────────────────────────────────────────
_ROOT = Path(__file__).resolve().parents[2]  # 项目根
FIXTURES_DIR = _ROOT / "tests" / "fixtures" / "v16_materials"


def _dir(subdir: str) -> Path:
    p = FIXTURES_DIR / subdir
    p.mkdir(parents=True, exist_ok=True)
    return p


# ──────────────────────────────────────────────────────────────────────────────
# txt 样本
# ──────────────────────────────────────────────────────────────────────────────
def make_txt_success() -> Path:
    p = _dir("txt") / "sample_success.txt"
    p.write_text(
        "这是一个测试文本文件。\n\n"
        "V16 R1 测试样本，用于验证 txt 解析链路（parse_txt_document）。\n"
        "包含多段正文，长度足够通过质量检查（quality_level=good）。\n"
        "段落2：Lorem ipsum dolor sit amet, consectetur adipiscing elit.\n",
        encoding="utf-8",
    )
    return p


def make_txt_empty() -> Path:
    p = _dir("failures") / "empty.txt"
    p.write_text("", encoding="utf-8")
    return p


def make_txt_low_quality() -> Path:
    p = _dir("failures") / "low_quality.txt"
    # 全是重复内容
    p.write_text("a\n" * 200, encoding="utf-8")
    return p


# ──────────────────────────────────────────────────────────────────────────────
# md 样本
# ──────────────────────────────────────────────────────────────────────────────
def make_md_success() -> Path:
    p = _dir("md") / "sample_success.md"
    p.write_text(
        "# V16 R1 Markdown 测试样本\n\n"
        "## 第一节：背景说明\n\n"
        "本文档用于验证 `parse_md_document` 工具的正向路径。\n\n"
        "## 第二节：内容\n\n"
        "- 要点一：支持 heading 级别识别\n"
        "- 要点二：支持列表解析\n"
        "- 要点三：保留结构信息\n\n"
        "## 结论\n\n"
        "parse_md_document 应产出 status=success，quality_level=good。\n",
        encoding="utf-8",
    )
    return p


# ──────────────────────────────────────────────────────────────────────────────
# docx 样本（用 python-docx 动态生成）
# ──────────────────────────────────────────────────────────────────────────────
def make_docx_success() -> Path:
    p = _dir("docx") / "sample_success.docx"
    try:
        import docx as _docx

        doc = _docx.Document()
        doc.add_heading("V16 R1 Word 文档测试", level=1)
        doc.add_heading("第一章：背景说明", level=2)
        doc.add_paragraph(
            "本 Word 文档用于验证 parse_docx 工具的正向路径。"
            "包含标题、段落、表格，预期 status=success，quality_level=good。"
        )
        doc.add_heading("第二章：表格测试", level=2)
        doc.add_paragraph("以下是一个示例表格：")
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "角色"
        table.cell(0, 2).text = "版本"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "施工执行者"
        table.cell(1, 2).text = "V16"
        table.cell(2, 0).text = "Bob"
        table.cell(2, 1).text = "质检负责人"
        table.cell(2, 2).text = "V16 R1"
        doc.add_paragraph("段落结尾：以上内容为测试数据。")
        doc.save(str(p))
    except ImportError:
        # python-docx 未安装时用最小 ZIP 骨架
        p.write_bytes(_minimal_docx_bytes())
    return p


def make_docx_empty() -> Path:
    p = _dir("failures") / "empty.docx"
    try:
        import docx as _docx

        doc = _docx.Document()
        doc.save(str(p))
    except ImportError:
        p.write_bytes(_minimal_docx_bytes())
    return p


def _minimal_docx_bytes() -> bytes:
    """生成一个极简合法 docx（空文档，ZIP 骨架）。"""
    try:
        import zipfile

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Override PartName="/word/document.xml"'
                ' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>",
            )
            z.writestr(
                "_rels/.rels",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"'
                ' Target="word/document.xml"/>'
                "</Relationships>",
            )
            z.writestr(
                "word/document.xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:body><w:p><w:r><w:t>Minimal docx content for V16 test.</w:t></w:r></w:p></w:body>"
                "</w:document>",
            )
        return buf.getvalue()
    except Exception:
        return b""


# ──────────────────────────────────────────────────────────────────────────────
# xlsx 样本（用 openpyxl 动态生成）
# ──────────────────────────────────────────────────────────────────────────────
def make_xlsx_success() -> Path:
    p = _dir("xlsx") / "sample_success.xlsx"
    try:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "销售数据"
        ws.append(["产品", "Q1销售额", "Q2销售额", "Q3销售额", "同比增长"])
        ws.append(["产品A", 10000, 12000, 15000, "50%"])
        ws.append(["产品B", 8000, 9500, 11000, "37.5%"])
        ws.append(["产品C", 5000, 6000, 7500, "50%"])
        ws2 = wb.create_sheet("人员名单")
        ws2.append(["姓名", "部门", "职位"])
        ws2.append(["张三", "研发部", "工程师"])
        ws2.append(["李四", "产品部", "产品经理"])
        wb.save(str(p))
    except ImportError:
        p.write_bytes(b"PK\x03\x04")  # ZIP magic bytes（不完整，仅占位）
    return p


def make_xlsx_empty() -> Path:
    p = _dir("failures") / "empty.xlsx"
    try:
        import openpyxl

        wb = openpyxl.Workbook()
        wb.active.title = "EmptySheet"
        wb.save(str(p))
    except ImportError:
        p.write_bytes(b"PK\x03\x04")
    return p


# ──────────────────────────────────────────────────────────────────────────────
# PDF 样本（用 reportlab 或 fpdf2 生成文本 PDF；无库则用嵌入字节串）
# ──────────────────────────────────────────────────────────────────────────────
def make_pdf_text_success() -> Path:
    p = _dir("pdf_text") / "sample_text.pdf"
    _make_text_pdf(p, title="V16 R1 文本 PDF 测试", pages=3)
    return p


def make_pdf_scanned_like() -> Path:
    """制造一个几乎无文本的 PDF（模拟扫描版，仅包含极少文字）。"""
    p = _dir("pdf_scanned") / "sample_scanned_like.pdf"
    _make_text_pdf(p, title="scan", pages=2, minimal_text=True)
    return p


def make_pdf_encrypted_placeholder() -> Path:
    """放置说明文件，真实加密 PDF 需手动放置。"""
    p = _dir("failures") / "README_encrypted_pdf.txt"
    p.write_text(
        "将加密的 PDF 文件重命名为 encrypted.pdf 放于此目录，\n"
        "测试 test_v16r1_parse_pdf.py::test_pdf_encrypted 会自动识别。\n",
        encoding="utf-8",
    )
    return p


def _make_text_pdf(p: Path, title: str, pages: int = 2, minimal_text: bool = False) -> None:
    """
    生成一个文本 PDF。

    优先用 PyMuPDF 直接创建（可保证可提取），否则用 fpdf2/reportlab，
    最后降级到内嵌最小 PDF 字节串（仅兜底）。
    minimal_text=True 时生成文本极少的页面（模拟扫描版密度）。
    """
    if _try_pymupdf_create(p, title, pages, minimal_text=minimal_text):
        return
    if not minimal_text:
        if _try_reportlab(p, title, pages):
            return
        if _try_fpdf2(p, title, pages):
            return
    p.write_bytes(_minimal_pdf_bytes(minimal_text=minimal_text))


def _try_pymupdf_create(
    p: Path, title: str, pages: int, minimal_text: bool = False
) -> bool:
    """用 PyMuPDF（fitz）创建 PDF 文件，保证文本可被 PyMuPDF 自身提取。"""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open()
        for i in range(pages):
            page = doc.new_page(width=595, height=842)  # A4
            if minimal_text:
                # 极少文字，模拟扫描版（<50 字符/页）
                page.insert_text((72, 72), ".", fontsize=12)
            else:
                text_block = (
                    f"{title} - Page {i + 1}\n"
                    f"V16 R1 test PDF content, page {i + 1} of {pages}.\n"
                    f"This is a text-based PDF for RAG pipeline testing.\n"
                    f"Lines include: page_number, block_index, extract_method.\n"
                    f"Source type: pdf. Tool: parse_pdf. MCP mode: mcp_compatible_adapter.\n"
                )
                page.insert_text((72, 72), text_block, fontsize=11)
        doc.save(str(p))
        doc.close()
        return True
    except ImportError:
        return False
    except Exception:
        return False


def _try_reportlab(p: Path, title: str, pages: int) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas as rl_canvas

        c = rl_canvas.Canvas(str(p), pagesize=A4)
        for i in range(pages):
            c.setFont("Helvetica", 12)
            c.drawString(72, 750, f"{title} - Page {i + 1}")
            c.drawString(72, 720, f"V16 R1 test PDF content, page {i + 1} of {pages}.")
            c.drawString(72, 690, "This is a text-based PDF for RAG pipeline testing.")
            c.drawString(72, 660, "Lines include: page_number, block_index, extract_method.")
            c.showPage()
        c.save()
        return True
    except ImportError:
        return False


def _try_fpdf2(p: Path, title: str, pages: int) -> bool:
    try:
        from fpdf import FPDF

        pdf = FPDF()
        for i in range(pages):
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            pdf.cell(0, 10, f"{title} - Page {i + 1}", ln=True)
            pdf.cell(0, 10, f"V16 R1 test PDF, page {i + 1}/{pages}.", ln=True)
            pdf.cell(0, 10, "Text PDF content for parse_pdf tool testing.", ln=True)
        pdf.output(str(p))
        return True
    except ImportError:
        return False


def _minimal_pdf_bytes(minimal_text: bool = False) -> bytes:
    """
    最终兜底：内嵌已知可用的最小 PDF 字节串（不依赖任何库）。
    minimal_text=True：极少文字（模拟扫描密度）。
    注：此 PDF 使用正确的 xref 偏移量，PyMuPDF 可以提取文本。
    """
    if minimal_text:
        # 1 个字符，模拟扫描版低文本密度
        content = b"BT /F1 12 Tf 72 720 Td (.) Tj ET"
    else:
        content = b"BT /F1 12 Tf 50 750 Td (V16 R1 test PDF - page_number block_index extract_method source_type pdf.) Tj ET"

    def _obj(n: int, body: bytes) -> bytes:
        return f"{n} 0 obj\n".encode() + body + b"\nendobj\n"

    hdr = b"%PDF-1.4\n"
    # obj 1: catalog
    o1 = _obj(1, b"<< /Type /Catalog /Pages 2 0 R >>")
    # obj 2: pages
    o2 = _obj(2, b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    # obj 3: page
    o3 = _obj(3, (
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
        b"   /Resources << /Font << /F1 5 0 R >> >>\n"
        b"   /Contents 4 0 R >>"
    ))
    # obj 4: content stream
    stream_len = len(content)
    o4 = (
        f"4 0 obj\n<< /Length {stream_len} >>\nstream\n".encode()
        + content
        + b"\nendstream\nendobj\n"
    )
    # obj 5: font
    o5 = _obj(5, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    # 计算字节偏移
    offsets = []
    body_parts = [o1, o2, o3, o4, o5]
    offset = len(hdr)
    for part in body_parts:
        offsets.append(offset)
        offset += len(part)

    # xref (6 entries: free + 5 objects)
    xref_offset = offset
    xref_lines = ["xref\n", "0 6\n", "0000000000 65535 f \n"]
    for off in offsets:
        xref_lines.append(f"{off:010d} 00000 n \n")
    xref = "".join(xref_lines).encode()
    trailer = (
        f"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    ).encode()

    return hdr + b"".join(body_parts) + xref + trailer


# ──────────────────────────────────────────────────────────────────────────────
# fixtures 总清单（供测试引用）
# ──────────────────────────────────────────────────────────────────────────────
FIXTURES_MANIFEST: dict[str, dict[str, Any]] = {
    "txt_success": {
        "factory": make_txt_success,
        "source_type": "txt",
        "expected_status": "success",
        "expected_quality_level": "good",
        "expected_error_code": "",
        "expected_pending_behavior": "pending 成功",
        "expected_commit_behavior": "commit 成功，可检索",
        "expected_retrieve_behavior": "retrieve 返回 chunk，metadata 含 source_type=txt",
    },
    "txt_empty": {
        "factory": make_txt_empty,
        "source_type": "txt",
        "expected_status": "failed",
        "expected_error_code": "empty_extracted_text",
        "expected_pending_behavior": "pending 失败态，不得 commit",
        "expected_commit_behavior": "commit 应失败",
        "expected_retrieve_behavior": "不可检索",
    },
    "md_success": {
        "factory": make_md_success,
        "source_type": "md",
        "expected_status": "success",
        "expected_quality_level": "good",
        "expected_error_code": "",
        "expected_pending_behavior": "pending 成功",
        "expected_commit_behavior": "commit 成功，可检索",
        "expected_retrieve_behavior": "retrieve 返回 chunk，metadata 含 source_type=md",
    },
    "docx_success": {
        "factory": make_docx_success,
        "source_type": "docx",
        "expected_status": "success",
        "expected_metadata": ["paragraph_count", "table_count", "extract_method"],
        "expected_error_code": "",
        "expected_pending_behavior": "pending 成功",
        "expected_commit_behavior": "commit 成功，可检索",
        "expected_retrieve_behavior": "retrieve 返回 chunk，metadata 含 source_type=docx",
    },
    "docx_empty": {
        "factory": make_docx_empty,
        "source_type": "docx",
        "expected_status": "failed",
        "expected_error_code": "docx_no_content|empty_extracted_text",
        "expected_pending_behavior": "pending 失败态",
        "expected_commit_behavior": "commit 应失败",
    },
    "xlsx_success": {
        "factory": make_xlsx_success,
        "source_type": "xlsx",
        "expected_status": "success",
        "expected_metadata": ["sheet_name", "column_names", "extract_method"],
        "expected_error_code": "",
        "expected_pending_behavior": "pending 成功",
        "expected_commit_behavior": "commit 成功，可检索",
        "expected_retrieve_behavior": "retrieve 返回 chunk，metadata 含 source_type=xlsx",
    },
    "xlsx_empty": {
        "factory": make_xlsx_empty,
        "source_type": "xlsx",
        "expected_status": "failed",
        "expected_error_code": "empty_sheet",
        "expected_pending_behavior": "pending 失败态",
        "expected_commit_behavior": "commit 应失败",
    },
    "pdf_text_success": {
        "factory": make_pdf_text_success,
        "source_type": "pdf",
        "expected_status": "success",
        "expected_metadata": ["page_count", "page_number", "is_scanned", "extract_method"],
        "expected_error_code": "",
        "expected_pending_behavior": "pending 成功",
        "expected_commit_behavior": "commit 成功，可检索",
        "expected_retrieve_behavior": "retrieve 返回 chunk，metadata 含 page_number",
    },
    "pdf_scanned_like": {
        "factory": make_pdf_scanned_like,
        "source_type": "pdf",
        "expected_status": "failed",
        "expected_error_code": "scanned_pdf_requires_ocr",
        "expected_metadata": ["is_scanned=True"],
        "expected_pending_behavior": "pending 失败态",
        "expected_commit_behavior": "commit 应失败",
    },
}


def ensure_fixtures() -> dict[str, Path]:
    """确保所有 fixtures 存在，返回 {name: path} 字典。"""
    result: dict[str, Path] = {}
    for name, spec in FIXTURES_MANIFEST.items():
        factory = spec["factory"]
        p = factory()
        result[name] = p
    return result


if __name__ == "__main__":
    paths = ensure_fixtures()
    for name, p in paths.items():
        print(f"  {name}: {p} ({'OK' if p.exists() else 'MISSING'})")
    print(f"\nfixtures 目录: {FIXTURES_DIR}")
